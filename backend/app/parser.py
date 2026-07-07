import io

import docx
import pdfplumber

from fastapi import HTTPException


def extract_text(filename: str, content: bytes) -> str:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return _extract_pdf(content)
    if lower.endswith(".docx"):
        return _extract_docx(content)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return content.decode("utf-8", errors="ignore")

    raise HTTPException(
        status_code=400,
        detail=f"Unsupported file type for '{filename}'. Upload a PDF, DOCX, or TXT file.",
    )


def _extract_pdf(content: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    text = "\n".join(text_parts).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="Could not extract any text from this PDF. It may be a scanned image without OCR text.",
        )
    return text


def _extract_docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs).strip()
    if not text:
        raise HTTPException(status_code=422, detail="Could not extract any text from this DOCX file.")
    return text
